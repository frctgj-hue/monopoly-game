# -*- coding: utf-8 -*-
"""
Универсальный скрипт для создания Word документов
Использование: python create_word.py <filename> <content_file>
"""
import sys
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document(filename, content):
    """Создает Word документ с форматированием"""
    doc = Document()
    
    # Парсим контент построчно
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Заголовок уровня 1
        if line.startswith('# '):
            doc.add_heading(line[2:], 0)
        # Заголовок уровня 2
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        # Заголовок уровня 3
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        # Маркированный список
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Нумерованный список
        elif line[0:2].replace('.', '').isdigit():
            doc.add_paragraph(line.split('. ', 1)[1] if '. ' in line else line, style='List Number')
        # Обычный параграф
        else:
            p = doc.add_paragraph(line)
    
    # Сохраняем
    desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
    output = os.path.join(desktop, f'{filename}.docx')
    doc.save(output)
    
    return output

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python create_word.py <filename> <content>')
        sys.exit(1)
    
    filename = sys.argv[1]
    content = sys.argv[2]
    
    try:
        output = create_word_document(filename, content)
        print(f'SUCCESS: {output}')
    except Exception as e:
        print(f'ERROR: {e}')
        sys.exit(1)
