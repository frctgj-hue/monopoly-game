# -*- coding: utf-8 -*-
from docx import Document
import os

doc = Document()
doc.add_heading('АБВГДЕЖЗИЙКЛМНОПРСТУФХЦЧШЩЭЮЯ', 0)
doc.add_paragraph('Александр Годунок')
doc.add_paragraph('Виды обращений')
doc.add_paragraph('Проверка всех букв алфавита')

desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
output = os.path.join(desktop, 'Тест_Script.docx')
doc.save(output)
print('Документ создан:', output)
